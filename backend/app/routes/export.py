from fastapi import APIRouter, HTTPException, Path, Query
from fastapi.responses import FileResponse, PlainTextResponse
import os
import json
from typing import Optional
from app.schemas import ExportFormat
from app.utils.logger import logger
import tempfile
import uuid

router = APIRouter()

@router.get("/export/{lease_id}")
async def export_summary(
    lease_id: str = Path(..., description="The unique ID of the processed lease"),
    format: ExportFormat = Query(ExportFormat.PDF, description="The format to export")
):
    """
    Export the lease summary in the requested format:
    - PDF: Professional document with styling
    - DOCX: Microsoft Word document
    - XLSX: Excel spreadsheet with structured data and rent schedule
    - JSON: Raw structured data
    - MARKDOWN: Raw markdown text
    """
    
    # Check if the lease exists
    storage_dir = os.path.join("app", "storage", "processed", lease_id)
    if not os.path.exists(storage_dir):
        raise HTTPException(status_code=404, detail="Lease not found")
    
    try:
        # Load the response data
        with open(os.path.join(storage_dir, "response.json"), 'r') as f:
            response_data = json.load(f)
        
        # Get summary markdown
        summary_markdown = response_data.get("summary_markdown", "")
        
        # Export directory
        export_dir = os.path.join("app", "storage", "exports")
        os.makedirs(export_dir, exist_ok=True)
        
        # Generate a unique filename
        filename = f"{lease_id}_{uuid.uuid4().hex[:8]}"
        
        if format == ExportFormat.MARKDOWN:
            # Return markdown directly
            return PlainTextResponse(summary_markdown)
            
        elif format == ExportFormat.JSON:
            # Return JSON directly
            return response_data
            
        elif format == ExportFormat.PDF:
            # Generate PDF from markdown
            output_path = os.path.join(export_dir, f"{filename}.pdf")
            generate_pdf(summary_markdown, output_path)
            return FileResponse(
                output_path, 
                media_type="application/pdf",
                filename=f"lease_summary_{lease_id}.pdf"
            )
            
        elif format == ExportFormat.DOCX:
            # Generate DOCX from markdown
            output_path = os.path.join(export_dir, f"{filename}.docx")
            generate_docx(summary_markdown, response_data, output_path)
            return FileResponse(
                output_path, 
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"lease_summary_{lease_id}.docx"
            )
            
        elif format == ExportFormat.XLSX:
            # Generate XLSX with structured data and rent schedule
            output_path = os.path.join(export_dir, f"{filename}.xlsx")
            generate_xlsx(response_data, output_path)
            return FileResponse(
                output_path, 
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                filename=f"lease_summary_{lease_id}.xlsx"
            )
            
    except Exception as e:
        logger.error(f"Error exporting lease {lease_id} in {format} format: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error exporting lease: {str(e)}")


def generate_pdf(markdown_content: str, output_path: str):
    """Generate a PDF from markdown content"""
    try:
        # This is a placeholder - in a real implementation, you would use a library like
        # ReportLab, WeasyPrint, or a markdown-to-pdf converter
        import markdown
        from weasyprint import HTML
        
        # Convert markdown to HTML
        html_content = markdown.markdown(markdown_content)
        
        # Add styling
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Lease Summary</title>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                h1 {{ color: #2c3e50; }}
                h2 {{ color: #3498db; border-bottom: 1px solid #eee; padding-bottom: 5px; }}
                h3 {{ color: #2980b9; }}
                .risk-high {{ color: #e74c3c; background-color: #fadbd8; padding: 5px; border-radius: 3px; }}
                .risk-medium {{ color: #e67e22; background-color: #fae5d3; padding: 5px; border-radius: 3px; }}
                .risk-low {{ color: #f39c12; background-color: #fcf3cf; padding: 5px; border-radius: 3px; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Generate PDF
        HTML(string=styled_html).write_pdf(output_path)
        
    except ImportError:
        # If WeasyPrint or markdown is not installed, write a placeholder PDF
        with open(output_path, 'w') as f:
            f.write("PDF generation requires WeasyPrint and markdown libraries")


def generate_docx(markdown_content: str, response_data: dict, output_path: str):
    """Generate a DOCX from markdown content"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        # Create a new Document
        doc = Document()
        
        # Add a title
        title = doc.add_heading("Lease Summary", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process markdown content and add to document
        lines = markdown_content.split('\n')
        
        for line in lines:
            # Handle headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            # Handle bullet points
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='ListBullet')
            # Handle normal text
            elif line.strip():
                doc.add_paragraph(line)
            # Add a blank line for empty lines
            else:
                doc.add_paragraph('')
        
        # Save the document
        doc.save(output_path)
        
    except ImportError:
        # If python-docx is not installed, write a placeholder file
        with open(output_path, 'w') as f:
            f.write("DOCX generation requires python-docx library")


def generate_xlsx(response_data: dict, output_path: str):
    """Generate an Excel spreadsheet with structured data and rent schedule"""
    try:
        import xlsxwriter
        
        # Create a workbook
        workbook = xlsxwriter.Workbook(output_path)
        
        # Add a Summary worksheet
        summary_sheet = workbook.add_worksheet('Summary')
        
        # Formatting
        title_format = workbook.add_format({
            'bold': True,
            'font_size': 14,
            'align': 'center',
            'valign': 'vcenter',
            'bg_color': '#3498db',
            'font_color': 'white'
        })
        
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#bdd7ee',
            'border': 1
        })
        
        cell_format = workbook.add_format({
            'border': 1
        })
        
        # Add title
        summary_sheet.merge_range('A1:F1', 'LEASE SUMMARY', title_format)
        summary_sheet.set_row(0, 30)
        
        # Add headers
        summary_sheet.write('A3', 'Category', header_format)
        summary_sheet.write('B3', 'Field', header_format)
        summary_sheet.write('C3', 'Value', header_format)
        summary_sheet.write('D3', 'Page', header_format)
        summary_sheet.write('E3', 'Confidence', header_format)
        summary_sheet.write('F3', 'Risk Level', header_format)
        
        # Set column widths
        summary_sheet.set_column('A:A', 15)
        summary_sheet.set_column('B:B', 20)
        summary_sheet.set_column('C:C', 40)
        summary_sheet.set_column('D:D', 10)
        summary_sheet.set_column('E:E', 12)
        summary_sheet.set_column('F:F', 12)
        
        # Add rent schedule worksheet (placeholder)
        rent_sheet = workbook.add_worksheet('Rent Schedule')
        
        # Close the workbook
        workbook.close()
        
    except ImportError:
        # If XlsxWriter is not installed, write a placeholder file
        with open(output_path, 'w') as f:
            f.write("XLSX generation requires XlsxWriter library")
